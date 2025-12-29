# src/rag/profile_rag.py

from __future__ import annotations

from pathlib import Path
from typing import List, Dict, Optional

import chromadb
from chromadb.utils import embedding_functions
from docx import Document  # pip install python-docx


# -------------------------------------------------------------------
# Directories
# -------------------------------------------------------------------
PROJECT_ROOT = Path(__file__).resolve().parents[2]
DATA_DIR = PROJECT_ROOT / "data"
UPLOADS_DIR = DATA_DIR / "uploads"
CHROMA_DIR = PROJECT_ROOT / ".chroma_profile"


# -------------------------------------------------------------------
# Globals for lazy init
# -------------------------------------------------------------------
_client: chromadb.PersistentClient | None = None
_collection: chromadb.Collection | None = None


# -------------------------------------------------------------------
# Chroma helpers
# -------------------------------------------------------------------
def _get_client() -> chromadb.PersistentClient:
    """
    Return a global PersistentClient for Chroma.
    """
    global _client
    if _client is None:
        CHROMA_DIR.mkdir(parents=True, exist_ok=True)
        _client = chromadb.PersistentClient(path=str(CHROMA_DIR))
    return _client


def _get_collection() -> chromadb.Collection:
    """
    Return (and lazily create) the Chroma collection
    for resume-based profile chunks.
    """
    global _collection
    if _collection is None:
        client = _get_client()
        _collection = client.get_or_create_collection(
            name="profile_resume_chunks",
            embedding_function=embedding_functions.DefaultEmbeddingFunction(),
        )
    return _collection


# -------------------------------------------------------------------
# Resume loading and chunking
# -------------------------------------------------------------------
def _get_latest_resume_path() -> Path:
    """
    Find the most recently modified .docx file in data/uploads.

    This supports filenames like:
      resume_20251224-xxxxx.docx, my_resume.docx, etc.
    """
    if not UPLOADS_DIR.exists():
        raise FileNotFoundError(
            f"Uploads directory not found at {UPLOADS_DIR}. "
            "Upload your resume from the UI first."
        )

    docx_files = list(UPLOADS_DIR.glob("*.docx"))
    if not docx_files:
        raise FileNotFoundError(
            f"No .docx resumes found in {UPLOADS_DIR}. "
            "Upload a resume (DOCX) from the UI."
        )

    # Pick newest by modification time
    latest = max(docx_files, key=lambda p: p.stat().st_mtime)
    return latest


def _load_latest_resume_text() -> str:
    """
    Load plain text from the latest uploaded resume .docx.
    """
    resume_path = _get_latest_resume_path()
    print(f"Using resume file for RAG: {resume_path}")

    doc = Document(str(resume_path))

    # Collect text from paragraphs
    lines: List[str] = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            lines.append(text)

    # Collect text from tables as well (common in resumes)[web:33]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = p.text.strip()
                    if text:
                        lines.append(text)

    return "\n".join(lines)


def chunk_resume(text: str, chunk_size: int = 1200, overlap: int = 200) -> List[str]:
    """
    Split resume text into overlapping fixed-size chunks (character-based).

    - chunk_size: max characters per chunk
    - overlap:    characters of overlap between consecutive chunks

    Fixed windows with modest overlap are recommended for RAG with Chroma.[web:25][web:50][web:54]
    """
    if not text:
        return []

    chunks: List[str] = []
    stride = max(chunk_size - overlap, 1)
    i = 0

    while i < len(text):
        chunk = text[i : i + chunk_size].strip()
        if chunk:
            chunks.append(chunk)
        i += stride

    return chunks


def _get_resume_chunks_from_latest_docx() -> List[Dict]:
    """
    Build chunk dicts from the FULL latest uploaded resume.

    Output format:
    [
      { "id": "resume_chunk_0", "text": "...", "section": "resume", "order": 0 },
      ...
    ]
    """
    full_text = _load_latest_resume_text()
    raw_chunks = chunk_resume(text=full_text, chunk_size=1200, overlap=200)

    chunks: List[Dict] = []
    for i, c in enumerate(raw_chunks):
        chunks.append(
            {
                "id": f"resume_chunk_{i}",
                "text": c,
                "section": "resume",
                "order": i,
            }
        )
    return chunks


# -------------------------------------------------------------------
# Index building / refreshing
# -------------------------------------------------------------------
def build_or_refresh_profile_index() -> None:
    """
    Rebuild the Chroma index from the latest uploaded resume file.
    """
    collection = _get_collection()
    print("Rebuilding profile index from resume chunks...")

    # ✅ CORRECTED: Get ALL existing IDs first, then delete
    current_count = collection.count()
    if current_count > 0:
        print(f"Found {current_count} existing documents. Deleting...")
        
        # Get all existing document IDs
        all_docs = collection.get(include=["metadatas"])
        existing_ids = all_docs['ids']
        
        if existing_ids:
            collection.delete(ids=existing_ids)
            print(f"Deleted {len(existing_ids)} documents.")
        else:
            print("No documents to delete.")
    else:
        print("Collection is empty - no delete needed.")

    # Add new chunks
    chunks: List[Dict] = _get_resume_chunks_from_latest_docx()
    print(f"Got {len(chunks)} chunks from latest uploaded resume.")

    if not chunks:
        print("No chunks found; leaving collection empty.")
        return

    ids = [c.get("id", f"resume_chunk_{i}") for i, c in enumerate(chunks)]
    texts = [c["text"] for c in chunks]
    metadatas = [
        {"section": c.get("section", "resume"), "order": c.get("order", i)}
        for i, c in enumerate(chunks)
    ]

    print(f"Adding {len(ids)} new documents...")
    collection.add(ids=ids, documents=texts, metadatas=metadatas)
    print(f"✅ Indexed {len(chunks)} resume-based profile chunks into Chroma.")

# -------------------------------------------------------------------
# Retrieval
# -------------------------------------------------------------------
def retrieve_relevant_chunks(
    query: str,
    top_k: int = 5,
) -> List[str]:
    """
    Given a query (e.g., job description or task description), return text
    of the top_k most relevant resume chunks from the latest uploaded resume.

    Assumes build_or_refresh_profile_index() has been called at least once
    after uploading/setting the resume.
    """
    collection = _get_collection()

    if collection.count() == 0:
        # Best-effort: try to build index now (in case it hasn't been built yet)
        try:
            build_or_refresh_profile_index()
        except FileNotFoundError as e:
            print(str(e))
            return []

    if collection.count() == 0:
        # Still empty → no resume / no chunks
        return []

    result = collection.query(
        query_texts=[query],
        n_results=top_k,
    )

    docs = result.get("documents", [[]])[0]
    return docs
